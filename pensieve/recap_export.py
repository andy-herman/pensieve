"""Render a recap payload to a Microsoft Word (.docx) document.

Takes the same recap dict produced by pensieve.recap.generate_recap and writes
a clean Connect-style document: a title, the reflection period, then one
heading per Connect goal with its accomplishment blocks (bold heading, narrative,
and a bold-labelled Impact line).
"""

from __future__ import annotations

import io
from typing import Any


def build_recap_docx(recap: dict[str, Any]) -> bytes:
    """Return the recap as .docx bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.add_heading("Connect Recap", level=0)
    period = recap.get("period_label") or ""
    if period:
        p = doc.add_paragraph()
        run = p.add_run(f"Reflection Period: {period}")
        run.italic = True

    scope = recap.get("scope", "all")
    considered = recap.get("memories_considered", 0)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Scope: {scope}  |  Tasks considered: {considered}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for section in recap.get("sections", []) or []:
        heading_text = section.get("short_name") or section.get("name") or "Goal"
        doc.add_heading(heading_text, level=1)
        full_name = section.get("name")
        if full_name and full_name != heading_text:
            sub = doc.add_paragraph()
            sub_run = sub.add_run(full_name)
            sub_run.italic = True
            sub_run.font.size = Pt(10)

        accomplishments = section.get("accomplishments", []) or []
        if not accomplishments:
            doc.add_paragraph("No accomplishment drafted for this goal.")
            continue
        for a in accomplishments:
            h = a.get("heading", "").strip()
            if h:
                hp = doc.add_paragraph()
                hp.add_run(h).bold = True
            narrative = a.get("narrative", "").strip()
            if narrative:
                doc.add_paragraph(narrative)
            impact = a.get("impact", "").strip()
            if impact:
                ip = doc.add_paragraph()
                ip.add_run("Impact: ").bold = True
                ip.add_run(impact)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
